from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Propuesta_Web_Reservas_Peluqueria.docx"


BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "2E7D32"
GOLD = "7A5A00"
RED = "9B1C1C"
TEXT = "1F2933"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=TEXT, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEXT)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_note(doc, title, body, color=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade_cell(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(TEXT)


def make_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=BLUE, size=9)
        if widths:
            table.columns[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.7)
            if widths:
                table.columns[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run("Propuesta de Pagina Web de Reservas")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(BLUE)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
r = subtitle.add_run("Peluqueria / Barberia - Prueba piloto")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(TEXT)

add_note(
    doc,
    "Resumen ejecutivo",
    "Se propone desarrollar una pagina web responsive para que los clientes puedan registrarse, consultar servicios y agendar citas desde el celular. El peluquero tendra un panel privado para gestionar citas, servicios, horarios, disponibilidad y pagos manuales. La primera etapa sera una prueba piloto para validar el uso real antes de invertir en funciones avanzadas.",
    LIGHT_BLUE,
)

add_section(doc, "1. Alcance de la primera version")
for item in [
    "Registro e inicio de sesion para clientes.",
    "Panel privado para el peluquero/administrador.",
    "Servicios con precio y duracion.",
    "Calendario de reservas con validacion para evitar doble reserva.",
    "Estados de cita: pendiente, confirmada, cancelada y finalizada.",
    "Estado del peluquero: disponible, ocupado, pausado o fuera de horario.",
    "Mapa de ubicacion del negocio.",
    "QR de pago o contacto.",
    "Registro manual de pagos.",
    "Diseno adaptable a celular, tablet y computador.",
]:
    add_bullet(doc, item)

add_section(doc, "2. Opciones de publicacion")
make_table(
    doc,
    ["Opcion", "Como quedaria", "Costo externo inicial", "Comentario"],
    [
        [
            "Servicios separados",
            "Vercel para frontend, Render para backend, Neon/Supabase para base de datos PostgreSQL y dominio aparte.",
            "Puede iniciar con planes gratis o de bajo costo. Dominio anual aparte.",
            "Mas economica para probar, pero se administra en varias plataformas.",
        ],
        [
            "Hostinger Negocio",
            "Pagina, backend Node.js, base de datos MySQL, dominio, SSL y correos dentro de Hostinger.",
            "COP $667.200 + impuestos por 48 meses + 3 meses gratis, segun captura del plan.",
            "Mas ordenada para el cliente: un solo proveedor y panel principal.",
        ],
    ],
    widths=[1.35, 2.45, 1.45, 1.25],
)

add_section(doc, "3. Opcion A: servicios separados")
doc.add_paragraph(
    "Esta opcion divide el proyecto en servicios especializados. Es util si se quiere iniciar con menor inversion y mantener PostgreSQL como base de datos."
)
make_table(
    doc,
    ["Servicio", "Uso", "Precio base consultado", "Observacion"],
    [
        ["Vercel", "Alojar el frontend.", "Plan Hobby: gratis.", "Ideal para iniciar la interfaz web."],
        ["Render", "Alojar el backend Node.js.", "Free con limite; Basic desde US$6/mes.", "Para estabilidad conviene plan pago."],
        ["Neon", "Base de datos PostgreSQL.", "Free: US$0/mes.", "Mantiene PostgreSQL como el proyecto actual."],
        ["Dominio", "Nombre de la pagina.", "Pago anual aparte.", "Ejemplo: www.nombrepeluqueria.com."],
    ],
    widths=[1.0, 1.85, 1.65, 2.0],
)
add_bullet(doc, "Ventaja principal: menor costo inicial y buena escalabilidad tecnica.")
add_bullet(doc, "Desventaja principal: hay mas cuentas, paneles y configuraciones por administrar.")

add_section(doc, "4. Opcion B: Hostinger Negocio")
doc.add_paragraph(
    "Esta opcion centraliza el proyecto en Hostinger. Es mas facil de explicar y administrar para el cliente porque todo queda en un solo proveedor."
)
make_table(
    doc,
    ["Concepto", "Detalle"],
    [
        ["Plan recomendado", "Hostinger Negocio."],
        ["Pago inicial", "COP $667.200 + impuestos por 48 meses, con 3 meses gratis."],
        ["Equivalente mensual promocional", "COP $13.900/mes, mostrado como referencia, pero el pago se realiza por el periodo completo."],
        ["Renovacion", "Despues del periodo inicial, aproximadamente COP $43.900/mes por 1 ano, segun la captura."],
        ["Incluye", "50 sitios web, 50 GB NVMe, dominio gratis 1 ano, 5 buzones por 1 ano, 5 apps Node.js, SSL, backups y CDN."],
        ["Base de datos", "Para que todo quede dentro de Hostinger se usaria MySQL, no PostgreSQL."],
    ],
    widths=[1.85, 4.65],
)
add_bullet(doc, "Ventaja principal: un solo proveedor, mas ordenado y mas facil para el cliente.")
add_bullet(doc, "Desventaja principal: pago inicial mas alto y cambio de PostgreSQL a MySQL.")

add_section(doc, "5. Pagos de las citas")
make_table(
    doc,
    ["Metodo", "Como funciona", "Costo/Complejidad", "Recomendacion"],
    [
        [
            "QR del negocio",
            "La pagina muestra el QR de Nequi, Daviplata, Bancolombia u otro medio. El pago llega directo al peluquero.",
            "Bajo. No requiere pasarela de pagos.",
            "Recomendado para la prueba piloto.",
        ],
        [
            "Wompi",
            "El cliente paga desde la pagina y el sistema puede confirmar pagos automaticamente.",
            "Mayor. Puede tener comisiones e integracion adicional.",
            "Dejar para una segunda fase.",
        ],
    ],
    widths=[1.0, 2.4, 1.45, 1.65],
)

add_section(doc, "6. Que pagaria el cliente")
for item in [
    "Desarrollo inicial de la pagina y del sistema de reservas: pago unico acordado con el desarrollador.",
    "Hosting o servicios de publicacion: Hostinger o servicios separados, segun la opcion elegida.",
    "Dominio: anual. En Hostinger Negocio puede estar incluido gratis el primer ano.",
    "Base de datos: incluida si se usa MySQL en Hostinger; externa si se usa Neon/Supabase.",
    "Mantenimiento tecnico mensual: opcional, si desea que el desarrollador administre la pagina.",
    "Integraciones futuras: Wompi, WhatsApp automatico, reportes avanzados o app movil, si se solicitan despues.",
]:
    add_bullet(doc, item)

add_section(doc, "7. Mantenimiento tecnico opcional")
doc.add_paragraph(
    "Si el cliente lo desea, el desarrollador puede quedar como administrador tecnico de la pagina. Este servicio seria mensual y separado del costo de hosting."
)
for item in [
    "Revisar que la pagina y el sistema sigan funcionando.",
    "Corregir errores tecnicos.",
    "Actualizar el sistema cuando sea necesario.",
    "Configurar dominio, hosting, variables de entorno y base de datos.",
    "Gestionar respaldos y revisar informacion importante.",
    "Implementar mejoras futuras solicitadas por el cliente.",
]:
    add_bullet(doc, item)

add_section(doc, "8. Recomendacion")
add_note(
    doc,
    "Recomendacion para iniciar",
    "Si el cliente quiere invertir menos al comienzo, conviene la opcion de servicios separados. Si el cliente prefiere orden, un solo proveedor y una gestion mas sencilla, conviene Hostinger Negocio. Para una presentacion formal y una administracion mas clara, se recomienda Hostinger Negocio; para validar con menor inversion, se recomienda la opcion separada.",
    LIGHT_GRAY,
)

add_section(doc, "Fuentes consultadas")
for source in [
    "Hostinger Web Hosting: plan Business/Negocio, Node.js apps, dominio, SSL, backups y precio promocional.",
    "Vercel Pricing: plan Hobby gratuito y Pro desde US$20/mes.",
    "Render Pricing: plan Free con limite y Basic desde US$6/mes.",
    "Neon Pricing: plan Free US$0/mes para PostgreSQL.",
    "Captura enviada del plan Hostinger Negocio en Colombia: COP $13.900/mes, pago total COP $667.200 + impuestos, renovacion COP $43.900/mes.",
]:
    add_bullet(doc, source)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Propuesta para prueba piloto - Pagina web de reservas")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string("6B7280")

doc.save(OUTPUT)
print(OUTPUT)
